import os
import sys
import argparse
import requests
import json
import re
import html
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def clean_html_tags(text):
    """
    Limpia las etiquetas HTML del texto y las convierte en formato adecuado.
    """
    # Convertir entidades HTML como &amp; a sus caracteres equivalentes
    text = html.unescape(text)
    
    # Etiquetas para eliminar completamente (como <br>)
    text = re.sub(r'<br\s*/?>', '\n', text)
    
    # Manejar etiquetas de centro
    text = re.sub(r'<center>(.*?)</center>', r'\1', text)
    
    # Manejar etiquetas de párrafo
    text = re.sub(r'<p>(.*?)</p>', r'\1\n', text)
    
    # Eliminar todas las demás etiquetas HTML
    text = re.sub(r'<[^>]*>', '', text)
    
    # Eliminar líneas en blanco múltiples
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    return text

def process_transcript_with_perplexity(transcript_path, api_key, custom_prompt=None, output_path=None):
    # Verificar que el archivo de transcripcion existe
    if not os.path.exists(transcript_path):
        print("Error: El archivo no existe")
        sys.exit(1)
    
    # Leer el contenido del archivo con mejor manejo de codificación
    try:
        # Intentar con UTF-8 primero
        with open(transcript_path, 'r', encoding='utf-8') as file:
            transcript_content = file.read()
        
        # Normalizar el texto para manejar caracteres especiales
        import unicodedata
        transcript_content = unicodedata.normalize('NFC', transcript_content)
        
        # Verificar que no haya caracteres problemáticos
        if any(ord(char) > 127 for char in transcript_content):
            print("Aviso: El archivo contiene caracteres especiales. Asegurando compatibilidad...")
        
    except UnicodeDecodeError:
        # Si falla con UTF-8, intentar con otra codificación (latin-1)
        try:
            with open(transcript_path, 'r', encoding='latin-1') as file:
                transcript_content = file.read()
            print("Aviso: Se utilizó codificación latin-1 para leer el archivo")
        except Exception as e:
            print(f"Error al leer el archivo con codificación alternativa: {str(e)}")
            sys.exit(1)
    except Exception as e:
        print(f"Error al leer el archivo: {str(e)}")
        sys.exit(1)    
    
    # Crear prompt si no se proporciona uno personalizado
    default_prompt = """Formatea esta transcripción como un acta formal profesional con las siguientes secciones (cuando estén disponibles):

1. TÍTULO: "ACTA DE ASAMBLEA" seguido del nombre de la organización si se menciona. Debe estar centrado.

2. DETALLES DE LA REUNIÓN: 
   - Fecha: [fecha]
   - Hora: [hora]
   - Lugar: [lugar]
   - Modalidad: [presencial/virtual]

3. ASISTENTES: 
   - Lista cada participante en una línea separada usando viñetas normales (no usar asteriscos ni símbolos especiales)
   - Formato: Nombre completo - Cargo (si se menciona)

4. ORDEN DEL DÍA: 
   - Lista numerada de temas a tratar
   - Un punto por línea

5. DESARROLLO: 
   - Resume los puntos principales discutidos
   - Usa párrafos para separar temas
   - Si hay acuerdos parciales, destácalos claramente

6. ACUERDOS: 
   - Lista cada acuerdo tomado en líneas separadas
   - Usa formato claro y conciso

7. TAREAS PENDIENTES: 
   - Indica responsable y fecha límite si se mencionan

8. CIERRE: 
   - Hora de finalización
   - Firmantes (si se mencionan)

INSTRUCCIONES IMPORTANTES DE FORMATO:
- NUNCA uses caracteres especiales como asteriscos (*), almohadillas (#), o símbolos extraños
- NUNCA uses formato Markdown de ningún tipo
- NUNCA uses etiquetas HTML como <center>, <br>, <p> o cualquier otra
- No uses negritas, cursivas ni subrayados - usa texto plano
- Para enfatizar texto, simplemente usa mayúsculas o describe la importancia
- Usa viñetas normales como "- " para listas (no uses asteriscos ni otros símbolos)
- Cuando necesites mostrar una etiqueta seguida de información (como "Fecha: 29 de abril"), usa el formato "Etiqueta: valor" sin negrita ni formato especial
- Separa claramente las secciones con líneas en blanco
- Si alguna sección no aparece en la transcripción, simplemente indica: "[Sección] no disponible en la transcripción"
- Usa un formato limpio y profesional, similar a un documento formal de una empresa
- No uses paréntesis para aclaraciones, mejor usa frases completas"""

    if not custom_prompt:
        custom_prompt = default_prompt
    
    # Preparar el prompt completo
    full_prompt = custom_prompt + "\n\n" + transcript_content
    
    # Preparar la llamada a la API de Perplexity
    headers = {
        "accept": "application/json",
        "content-type": "application/json",
        "authorization": "Bearer " + api_key
    }
    
    payload = {
        "model": "sonar",
        "messages": [
            {
                "role": "system",
                "content": "Eres un asistente experto en formatear transcripciones de reuniones en actas formales con un formato profesional y claro. IMPORTANTE: No uses formato Markdown ni etiquetas HTML. Específicamente, no uses asteriscos (*) para negrita, ni etiquetas <br>, <center>, o cualquier otra etiqueta HTML. Usa únicamente texto plano con espacios y saltos de línea normales."
            },
            {
                "role": "user",
                "content": full_prompt
            }
        ],
        "temperature": 0.5,
        "max_tokens": 4000
    }
    
    # URL de la API de Perplexity
    api_url = "https://api.perplexity.ai/chat/completions"
    
    # Realizar la llamada a la API
    print("Llamando a la API de Perplexity...")
    print(f"Usando clave API: {api_key[:5]}...{api_key[-5:]}")
    try:
        response = requests.post(api_url, headers=headers, json=payload)
        print(f"Código de estado: {response.status_code}")
        print(f"Respuesta: {response.text[:500]}")
        response.raise_for_status()
    except Exception as e:
        print(f"Error al llamar a la API de Perplexity: {type(e).__name__} - {str(e)}")
        sys.exit(1)
    
    # Procesar la respuesta
    try:
        result = response.json()
        formatted_transcript = result['choices'][0]['message']['content']
        
        # Limpiar cualquier resto de formato markdown, especialmente asteriscos
        # Reemplazar los patrones de texto en negrita de markdown
        bold_pattern = re.compile(r'\*\*(.*?)\*\*')
        italic_pattern = re.compile(r'\*(.*?)\*')
        
        # Guardamos los pares de etiqueta-valor para formatearlos después
        label_value_pairs = []
        label_pattern = re.compile(r'(\w+[\s\w]*?):\s+(.*)')
        
        # Eliminar asteriscos de negrita (ignorando los que forman parte de viñetas)
        formatted_transcript = bold_pattern.sub(r'\1', formatted_transcript)
        
        # Eliminar asteriscos de cursiva
        formatted_transcript = italic_pattern.sub(r'\1', formatted_transcript)
        
        # Normalizar las viñetas para usar un formato consistente
        formatted_transcript = re.sub(r'^\*\s+', '- ', formatted_transcript, flags=re.MULTILINE)
        
        # Limpiar etiquetas HTML
        formatted_transcript = clean_html_tags(formatted_transcript)
        
        print("Texto procesado y limpiado de formatos Markdown y etiquetas HTML")
        
    except Exception as e:
        print(f"Error al procesar la respuesta de la API: {str(e)}")
        sys.exit(1)
    
    # Determinar la ruta de salida si no se proporciona
    if not output_path:
        base_dir = os.path.dirname(transcript_path)
        base_name = os.path.splitext(os.path.basename(transcript_path))[0]
        output_path = os.path.join(base_dir, base_name + "_acta_formatada.txt")
    
    # Guardar la respuesta en un archivo de texto plano con formato enriquecido
    try:
        # Normalizar el texto
        import unicodedata
        normalized_transcript = unicodedata.normalize('NFC', formatted_transcript)
        
        # Guardar en formato TXT como antes
        with open(output_path, 'wb') as file:
            file.write(b'\xef\xbb\xbf')  # UTF-8 BOM
            file.write(normalized_transcript.encode('utf-8', errors='strict'))
        
        print(f"Acta formatada guardada en: {output_path}")
        
        try:
            # También crear una versión Word del documento
            docx_path = output_path.replace('.txt', '.docx')
            print(f"Generando documento Word en: {docx_path}")
            
            # Crear un nuevo documento Word
            doc = Document()
            
            # Configurar estilos del documento
            styles = doc.styles
            
            # Estilo para título principal
            style_title = styles['Title']
            style_title.font.name = 'Calibri'
            style_title.font.size = Pt(16)
            style_title.font.bold = True
            style_title.font.color.rgb = RGBColor(0, 0, 102)  # Azul oscuro
            
            # Estilo para encabezados de sección
            style_heading = styles['Heading 1']
            style_heading.font.name = 'Calibri'
            style_heading.font.size = Pt(14)
            style_heading.font.bold = True
            style_heading.font.color.rgb = RGBColor(0, 51, 102)  # Azul oscuro
            
            # Estilo para texto normal
            style_normal = styles['Normal']
            style_normal.font.name = 'Calibri'
            style_normal.font.size = Pt(11)
            
            # Función para agregar línea horizontal
            def add_horizontal_line(paragraph):
                p = paragraph._p  # _p es el elemento xml paragraph
                pPr = p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                pPr.append(pBdr)
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), '000000')
                pBdr.append(bottom)
            
            # Procesar el texto formateado línea por línea para aplicar estilos
            lines = normalized_transcript.split('\n')
            
            # Buscar título principal y secciones
            is_first_line = True
            in_list = False
            current_section = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    # Línea vacía: agregar un párrafo en blanco
                    doc.add_paragraph()
                    in_list = False
                    continue
                
                # Procesar según el contenido de la línea
                if is_first_line or line.upper().startswith("ACTA DE") or "ACTA " in line.upper() and len(line) < 50:
                    # Es un título principal
                    if not is_first_line:  # Si no es la primera línea
                        p = doc.add_paragraph()
                        add_horizontal_line(p)
                    title = doc.add_paragraph(line)
                    title.style = 'Title'
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    is_first_line = False
                    
                    # Agregar línea horizontal después del título
                    p = doc.add_paragraph()
                    add_horizontal_line(p)
                    continue
                
                # Detectar encabezados de sección
                section_keywords = ["FECHA", "DETALLES DE LA REUNIÓN", "ASISTENTES", "PARTICIPANTES", 
                                    "QUÓRUM", "ORDEN DEL DÍA", "DESARROLLO", "ACUERDOS", 
                                    "TAREAS PENDIENTES", "CIERRE", "INTRODUCCIÓN"]
                
                is_section_header = False
                for keyword in section_keywords:
                    if keyword in line.upper() and len(line) < 60:
                        # Es un encabezado de sección
                        heading = doc.add_paragraph(line)
                        heading.style = 'Heading 1'
                        current_section = keyword
                        is_section_header = True
                        in_list = False
                        break
                
                if is_section_header:
                    continue
                
                # Identificar líneas con formato especial "Etiqueta: Valor"
                # Estas suelen tener estructura como "Fecha:", "Hora:", etc.
                label_match = re.match(r'^([^:]+):\s*(.*)', line)
                if label_match and len(label_match.group(1)) < 25:
                    label_part = label_match.group(1).strip()
                    value_part = label_match.group(2).strip()
                    
                    p = doc.add_paragraph()
                    # Aplicar negrita a la etiqueta
                    p.add_run(label_part + ":").bold = True
                    # Añadir el valor en texto normal
                    p.add_run(" " + value_part)
                    continue
                
                # Identificar listas y viñetas
                if line.strip().startswith("- ") or line.strip().startswith("• "):
                    # Es un elemento de lista
                    text = line[2:].strip()
                    p = doc.add_paragraph(text, style='List Bullet')
                    p_format = p.paragraph_format
                    p_format.left_indent = Inches(0.25)
                    in_list = True
                elif line.strip().startswith("1.") or line.strip().startswith("1)") or \
                    (line[0:1].isdigit() and line[1:2] in [".", ")", ":"] and len(line) > 3):
                    # Es una lista numerada
                    text = line[line.find(" ")+1:].strip()
                    # Extraer el número
                    num = line[:line.find(" ")].rstrip(".):,")
                    p = doc.add_paragraph(text, style='List Number')
                    p_format = p.paragraph_format
                    p_format.left_indent = Inches(0.25)
                    in_list = True
                elif in_list and (line.startswith("  ") or line.startswith("\t")):
                    # Continuación de un elemento de lista
                    p = doc.add_paragraph(line.strip())
                    p_format = p.paragraph_format
                    p_format.left_indent = Inches(0.5)
                    p.style = 'List Bullet 2'
                else:
                    # Texto normal
                    p = doc.add_paragraph(line)
                    p.style = 'Normal'
                    in_list = False
            
            # Guardar el documento Word
            doc.save(docx_path)
            print(f"Documento Word guardado en: {docx_path}")
        except ImportError as ie:
            print(f"AVISO: No se pudo crear el documento Word porque falta la biblioteca python-docx: {str(ie)}")
            print("Por favor, instale python-docx con: pip install python-docx")
        except Exception as e:
            print(f"Error al crear documento Word: {str(e)}")
            print("Continuando con solo la versión de texto plano...")
    
    except Exception as e:
        print(f"Error al guardar el archivo de salida: {str(e)}")
        sys.exit(1)
        
    return output_path

def main():
    # Configurar el parser de argumentos
    parser = argparse.ArgumentParser(description='Procesa una transcripcion con la API de Perplexity')
    parser.add_argument('transcript_path', help='Ruta del archivo de transcripcion')
    parser.add_argument('--api_key', help='API Key de Perplexity')
    parser.add_argument('--prompt', help='Prompt personalizado para la IA')
    parser.add_argument('--output', help='Ruta para guardar el resultado')
    
    args = parser.parse_args()
    
    # Verificar API Key
    api_key = args.api_key or os.environ.get('PERPLEXITY_API_KEY')
    if not api_key:
        print("Error: Se requiere API Key de Perplexity")
        sys.exit(1)
    
    # Procesar la transcripcion
    try:
        output_file = process_transcript_with_perplexity(
            args.transcript_path, 
            api_key,
            args.prompt,
            args.output
        )
        print("Procesamiento completado con exito")
        sys.exit(0)
    except Exception as e:
        print(f"Error durante el procesamiento: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()